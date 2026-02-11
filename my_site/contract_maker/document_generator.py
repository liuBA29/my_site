# Генерация договора подряда и акта выполненных работ для самозанятой (исполнитель)
# Адаптировано из my_personal_contract_maker для Django

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from num2words import num2words

from .generator_config import (
    FULL_NAME,
    SHORT_NAME,
    INN,
    ADDRESS,
    BASIS,
    BANK_DATA,
    OUTPUT_DIR,
    TEMPLATES_DIR,
)


def _style_doc(doc):
    """Общие стиль и поля для документов."""
    style = doc.styles["Normal"]
    style.font.name = "Courier New"
    style.font.size = Pt(10)
    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.5)


def _replace_in_doc(doc, replacements):
    """Подставляет в документе все {{KEY}} на значения из словаря replacements."""
    for p in doc.paragraphs:
        text = p.text
        for key, value in replacements.items():
            text = text.replace("{{" + key + "}}", str(value))
        if text != p.text:
            p.clear()
            p.add_run(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text
                    for key, value in replacements.items():
                        text = text.replace("{{" + key + "}}", str(value))
                    if text != p.text:
                        p.clear()
                        p.add_run(text)


def _customer_requisites(customer):
    parts = [customer.get("address", "")]
    if customer.get("unp"):
        parts.append(f"УНП: {customer['unp']}")
    if customer.get("okpo"):
        parts.append(f"ОКПО: {customer['okpo']}")
    if customer.get("iban"):
        parts.append(f"IBAN: {customer['iban']}")
    return "\n".join(p for p in parts if p).strip()


def _executor_requisites():
    return f"ИНН: {INN}\n{ADDRESS}\n{BANK_DATA}".strip() if BANK_DATA else f"ИНН: {INN}\n{ADDRESS}"


def _add_requisites_and_signatures(doc, customer, executor_short_name, customer_short_name, is_individual=False):
    """Реквизиты сторон и подписи (общее для договора и акта). is_individual: заказчик — физлицо (без реквизитов юрлица)."""
    p = doc.add_paragraph("Реквизиты сторон:", style="Normal")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.runs[0].font.bold = True
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Pt(200)
    table.columns[1].width = Pt(200)
    row_cells = table.rows[0].cells

    p_ex = row_cells[0].paragraphs[0]
    r1 = p_ex.add_run(f"{FULL_NAME}\n")
    r1.font.size = Pt(9)
    r1.bold = True
    r2 = p_ex.add_run(f"ИНН: {INN}\n{ADDRESS}\n{BANK_DATA}".strip())
    r2.font.size = Pt(9)

    p_cl = row_cells[1].paragraphs[0]
    r1 = p_cl.add_run(f"{customer.get('table_name', customer.get('org_name', ''))}\n")
    r1.font.size = Pt(9)
    r1.bold = True
    if is_individual:
        details = customer.get("address", "")
    else:
        details = customer.get("address", "")
        if customer.get("unp"):
            details += f"\nУНП: {customer['unp']}"
        if customer.get("okpo"):
            details += f", ОКПО: {customer['okpo']}"
        if customer.get("iban"):
            details += f"\nIBAN: {customer['iban']}"
    r2 = p_cl.add_run(details.strip())
    r2.font.size = Pt(9)

    table.columns[0].width = Pt(150)
    table.columns[1].width = Pt(350)

    sig = doc.add_table(rows=1, cols=2)
    sig.columns[0].width = Pt(150)
    sig.columns[1].width = Pt(450)
    sig.cell(0, 0).text = f"Исполнитель_________({executor_short_name})"
    sig.cell(0, 0).paragraphs[0].runs[0].font.size = Pt(9)
    sig.cell(0, 1).text = f"Заказчик_________({customer_short_name})"
    sig.cell(0, 1).paragraphs[0].runs[0].font.size = Pt(9)


CONTRACT_TEMPLATE_NAME = "dogovor-podryada.docx"
ACT_TEMPLATE_NAME = "akt.docx"


def generate_contract(
    customer,
    work_list,
    contract_number,
    doc_date,
    location,
    total_cost,
    payment_percent=100,
    is_prepay=True,
    save_dir=None,
):
    """Генерирует договор подряда (docx)."""
    template_path = TEMPLATES_DIR / CONTRACT_TEMPLATE_NAME
    if template_path.exists():
        return _generate_contract_from_template(
            template_path, customer, work_list, contract_number, doc_date,
            location, total_cost, payment_percent, is_prepay, save_dir,
        )
    return _generate_contract_builtin(
        customer, work_list, contract_number, doc_date, location, total_cost,
        payment_percent, is_prepay, save_dir,
    )


def _generate_contract_from_template(
    template_path,
    customer,
    work_list,
    contract_number,
    doc_date,
    location,
    total_cost,
    payment_percent,
    is_prepay,
    save_dir,
):
    rubles = int(total_cost)
    sum_words = num2words(rubles, lang="ru", to="cardinal").capitalize()
    if is_prepay and payment_percent == 100:
        payment_terms = (
            "Заказчик производит 100% предоплату работ в течение 5(пяти) рабочих дней "
            "с момента подписания сторонами настоящего договора."
        )
    elif is_prepay and payment_percent:
        payment_terms = (
            f"Заказчик производит {payment_percent}% предоплату работ в течение 5(пяти) рабочих дней "
            "с момента подписания сторонами настоящего договора. Оставшаяся часть — в течение 5(пяти) банковских дней "
            "после подписания сторонами Акта сдачи-приемки работ, являющегося Приложением 2 к настоящему договору. "
            "Заказчик по своему усмотрению может произвести полную предоплату работ."
        )
    else:
        payment_terms = (
            "Заказчик производит оплату Работ в течение 5(пяти) банковских дней после подписания "
            "сторонами Акта сдачи-приемки работ, являющегося Приложением 2 к настоящему договору. "
            "Заказчик, по своему усмотрению, может произвести полную или частичную предоплату работ."
        )
    if is_prepay:
        completion_terms = (
            f"Окончание Работ: в течение 7(семи) рабочих дней с момента поступления на счёт Исполнителя "
            f"{payment_percent or 100}% предоплаты."
        )
    else:
        completion_terms = (
            "Окончание Работ: в течение 7(семи) рабочих дней с момента подписания сторонами настоящего договора. "
            "Сроки выполнения Работ могут быть продлены по согласованию сторон."
        )
    rep_pos = customer.get("rep_position", "")
    rep_name = customer.get("rep_name", "")
    customer_rep = f"{rep_pos} {rep_name}".strip()
    replacements = {
        "CONTRACT_NUMBER": contract_number,
        "DOC_DATE": doc_date,
        "LOCATION": location,
        "CUSTOMER_ORG": customer.get("org_name", ""),
        "CUSTOMER_REP": customer_rep,
        "CUSTOMER_BASIS": customer.get("basis", ""),
        "EXECUTOR_NAME": FULL_NAME,
        "EXECUTOR_BASIS": BASIS,
        "WORK_LIST": "\n".join(f"{i}. {w}" for i, w in enumerate(work_list, 1)),
        "TOTAL_COST": f"{total_cost:.2f}",
        "TOTAL_COST_WORDS": sum_words,
        "PAYMENT_TERMS": payment_terms,
        "COMPLETION_TERMS": completion_terms,
        "EXECUTOR_REQUISITES": _executor_requisites(),
        "CUSTOMER_REQUISITES": _customer_requisites(customer),
        "SIGNATURE_EXECUTOR": SHORT_NAME,
        "SIGNATURE_CUSTOMER": customer.get("short_name", customer.get("rep_name", "")),
    }
    doc = Document(template_path)
    _replace_in_doc(doc, replacements)
    save_dir = Path(save_dir or OUTPUT_DIR)
    save_dir.mkdir(parents=True, exist_ok=True)
    out_path = save_dir / f"contract_{contract_number}.docx"
    doc.save(out_path)
    return str(out_path)


def _generate_contract_builtin(
    customer,
    work_list,
    contract_number,
    doc_date,
    location,
    total_cost,
    payment_percent=100,
    is_prepay=True,
    save_dir=None,
):
    doc = Document()
    _style_doc(doc)

    p = doc.add_paragraph(f"ДОГОВОР № {contract_number}", style="Normal")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.bold = True

    tbl = doc.add_table(rows=1, cols=2)
    tbl.cell(0, 0).text = "г. Минск"
    tbl.cell(0, 0).paragraphs[0].runs[0].font.size = Pt(10)
    tbl.cell(0, 1).text = doc_date
    tbl.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    tbl.cell(0, 1).paragraphs[0].runs[0].font.size = Pt(10)

    org = customer.get("org_name", "")
    is_individual = customer.get("client_type") == "individual"
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    if is_individual:
        p.add_run(
            f"1. Гражданин {org}, именуемый в дальнейшем ЗАКАЗЧИК, с одной стороны, и {FULL_NAME}, "
            f"действующая на основании {BASIS}, именуемая в дальнейшем ИСПОЛНИТЕЛЬ, "
            f"с другой стороны, заключили настоящий договор о нижеследующем:"
        ).font.size = Pt(10)
    else:
        rep_pos = customer.get("rep_position", "")
        rep_name = customer.get("rep_name", "")
        cust_basis = customer.get("basis", "")
        p.add_run(
            f"1. {org}, в лице {rep_pos} {rep_name}, действующего на основании {cust_basis}, "
            f"именуемое в дальнейшем ЗАКАЗЧИК, с одной стороны, и {FULL_NAME}, "
            f"действующая на основании {BASIS}, именуемая в дальнейшем ИСПОЛНИТЕЛЬ, "
            f"с другой стороны, заключили настоящий договор о нижеследующем:"
        ).font.size = Pt(10)
    p.paragraph_format.space_after = Pt(0)

    doc.add_paragraph(
        "1. ЗАКАЗЧИК поручает, а ИСПОЛНИТЕЛЬ принимает на себя выполнение следующих работ:",
        style="Normal",
    ).paragraph_format.space_after = Pt(0)
    for idx, work in enumerate(work_list, 1):
        doc.add_paragraph(f"1.{idx}. {work}.", style="Normal").paragraph_format.space_after = Pt(0)

    doc.add_paragraph(f"2. Место проведения работ: {location}.", style="Normal").paragraph_format.space_after = Pt(0)

    rubles = int(total_cost)
    sum_words = num2words(rubles, lang="ru", to="cardinal")
    doc.add_paragraph(
        f"3. Стоимость работ, согласно Приложению 1 к настоящему договору, составляет {total_cost:.2f} "
        f"({sum_words.capitalize()} белорусских рублей 00 коп.). Без НДС.",
        style="Normal",
    ).paragraph_format.space_after = Pt(0)

    if is_prepay and payment_percent == 100:
        pay_text = (
            "4. Условия оплаты: Заказчик производит 100% предоплату работ в течение 5(пяти) рабочих дней "
            "с момента подписания сторонами настоящего договора."
        )
    elif is_prepay and payment_percent:
        pay_text = (
            f"4. Условия оплаты: Заказчик производит {payment_percent}% предоплату работ в течение 5(пяти) рабочих дней "
            "с момента подписания сторонами настоящего договора. Оставшаяся часть — в течение 5(пяти) банковских дней "
            "после подписания сторонами Акта сдачи-приемки работ, являющегося Приложением 2 к настоящему договору. "
            "Заказчик по своему усмотрению может произвести полную предоплату работ."
        )
    else:
        pay_text = (
            "4. Условия оплаты: Заказчик производит оплату Работ в течение 5(пяти) банковских дней после подписания "
            "сторонами Акта сдачи-приемки работ, являющегося Приложением 2 к настоящему договору. "
            "Заказчик, по своему усмотрению, может произвести полную или частичную предоплату работ."
        )
    doc.add_paragraph(pay_text, style="Normal").paragraph_format.space_after = Pt(0)

    if is_prepay:
        completion_text = (
            f"5. Окончание Работ: в течение 7(семи) рабочих дней с момента поступления на счёт Исполнителя "
            f"{payment_percent or 100}% предоплаты."
        )
    else:
        completion_text = (
            "5. Окончание Работ: в течение 7(семи) рабочих дней с момента подписания сторонами настоящего договора. "
            "Сроки выполнения Работ могут быть продлены по согласованию сторон."
        )
    doc.add_paragraph(completion_text, style="Normal").paragraph_format.space_after = Pt(0)
    doc.add_paragraph(
        "6. Заказчик в день получения акта сдачи-приемки работ обязан вернуть Исполнителю подписанный акт "
        "или мотивированный отказ от приемки.",
        style="Normal",
    ).paragraph_format.space_after = Pt(0)
    doc.add_paragraph(
        "7. По всем остальным вопросам стороны руководствуются действующим законодательством Республики Беларусь.",
        style="Normal",
    ).paragraph_format.space_after = Pt(0)

    _add_requisites_and_signatures(
        doc,
        customer,
        executor_short_name=SHORT_NAME,
        customer_short_name=customer.get("short_name", customer.get("rep_name", customer.get("org_name", ""))),
        is_individual=is_individual,
    )

    doc.add_paragraph()
    p = doc.add_paragraph(f"Приложение 1 к ДОГОВОРУ № {contract_number} от {doc_date}", style="Normal")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(11)
    doc.add_paragraph("Протокол согласования договорной цены", style="Normal").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.columns[0].width = Pt(50)
    table.columns[1].width = Pt(350)
    table.columns[2].width = Pt(100)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "№", "Наименование работ", "Стоимость, руб."
    for c in hdr:
        c.paragraphs[0].runs[0].font.bold = True
        c.paragraphs[0].runs[0].font.size = Pt(10)
        c.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    num_works = len(work_list)
    price_per = total_cost / num_works if num_works else 0
    for idx, work in enumerate(work_list, 1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = work
        row[2].text = f"{price_per:.2f}"
        row[2].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    row = table.add_row().cells
    row[0].text = "Итого:"
    row[1].merge(row[2])
    row[1].text = f"{total_cost:.2f}"
    row[1].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    for r in row[1].paragraphs[0].runs:
        r.font.bold = True
        r.font.size = Pt(10)

    p2 = doc.add_paragraph(style="Normal")
    p2.add_run("Итого стоимость работ: ").font.bold = True
    p2.add_run(f"{sum_words.capitalize()} белорусских рублей 00 копеек.\n").font.size = Pt(10)
    p2.add_run("Без НДС.").font.size = Pt(10)

    sig2 = doc.add_table(rows=1, cols=2)
    sig2.cell(0, 0).text = f"Исполнитель_________({SHORT_NAME})"
    sig2.cell(0, 0).paragraphs[0].runs[0].font.size = Pt(9)
    sig2.cell(0, 1).text = f"Заказчик_________({customer.get('short_name', '')})"
    sig2.cell(0, 1).paragraphs[0].runs[0].font.size = Pt(9)

    save_dir = Path(save_dir or OUTPUT_DIR)
    save_dir.mkdir(parents=True, exist_ok=True)
    file_path = save_dir / f"contract_{contract_number}.docx"
    doc.save(file_path)
    return str(file_path)


def generate_act(
    customer,
    work_list,
    contract_number,
    doc_date,
    act_date,
    total_cost,
    save_dir=None,
):
    """Генерирует акт сдачи-приемки работ (docx)."""
    template_path = TEMPLATES_DIR / ACT_TEMPLATE_NAME
    if template_path.exists():
        return _generate_act_from_template(
            template_path, customer, work_list, contract_number, doc_date, act_date, total_cost, save_dir,
        )
    return _generate_act_builtin(
        customer, work_list, contract_number, doc_date, act_date, total_cost, save_dir,
    )


def _generate_act_from_template(
    template_path, customer, work_list, contract_number, doc_date, act_date, total_cost, save_dir,
):
    rubles = int(total_cost)
    sum_words = num2words(rubles, lang="ru", to="cardinal").capitalize()
    replacements = {
        "CONTRACT_NUMBER": contract_number,
        "DOC_DATE": doc_date,
        "ACT_DATE": act_date,
        "CUSTOMER_ORG": customer.get("org_name", ""),
        "CUSTOMER_REP": f"{customer.get('rep_position', '')} {customer.get('rep_name', '')}".strip(),
        "CUSTOMER_BASIS": customer.get("basis", ""),
        "EXECUTOR_NAME": FULL_NAME,
        "EXECUTOR_BASIS": BASIS,
        "WORK_LIST": "\n".join(f"{i}. {w}" for i, w in enumerate(work_list, 1)),
        "TOTAL_COST": f"{total_cost:.2f}",
        "TOTAL_COST_WORDS": sum_words,
        "EXECUTOR_REQUISITES": _executor_requisites(),
        "CUSTOMER_REQUISITES": _customer_requisites(customer),
        "SIGNATURE_EXECUTOR": SHORT_NAME,
        "SIGNATURE_CUSTOMER": customer.get("short_name", customer.get("rep_name", "")),
    }
    doc = Document(template_path)
    _replace_in_doc(doc, replacements)
    save_dir = Path(save_dir or OUTPUT_DIR)
    save_dir.mkdir(parents=True, exist_ok=True)
    out_path = save_dir / f"act_{contract_number}.docx"
    doc.save(out_path)
    return str(out_path)


def _generate_act_builtin(
    customer,
    work_list,
    contract_number,
    doc_date,
    act_date,
    total_cost,
    save_dir=None,
):
    doc = Document()
    _style_doc(doc)

    p = doc.add_paragraph(
        f"Приложение 2 к ДОГОВОРУ № {contract_number}  (от {doc_date} года)",
        style="Normal",
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.runs[0].font.size = Pt(12)
    p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph("Акт сдачи-приемки работ", style="Normal")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.bold = True

    tbl = doc.add_table(rows=1, cols=2)
    tbl.cell(0, 0).text = "г. Минск"
    tbl.cell(0, 0).paragraphs[0].runs[0].font.size = Pt(11)
    tbl.cell(0, 1).text = act_date
    tbl.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    tbl.cell(0, 1).paragraphs[0].runs[0].font.size = Pt(11)

    org = customer.get("org_name", "")
    is_individual = customer.get("client_type") == "individual"
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    if is_individual:
        p.add_run(
            f"1. Гражданин {org}, именуемый в дальнейшем ЗАКАЗЧИК, с одной стороны, и {FULL_NAME}, "
            f"действующая на основании {BASIS}, именуемая в дальнейшем ИСПОЛНИТЕЛЬ, с другой стороны, "
            f"составили настоящий Акт о том, что выполненные Исполнителем работы:"
        ).font.size = Pt(11)
    else:
        rep_pos = customer.get("rep_position", "")
        rep_name = customer.get("rep_name", "")
        cust_basis = customer.get("basis", "")
        p.add_run(
            f"1. {org}, в лице {rep_pos} {rep_name}, действующего на основании {cust_basis}, именуемое в дальнейшем ЗАКАЗЧИК, "
            f"с одной стороны, и {FULL_NAME}, действующая на основании {BASIS}, именуемая в дальнейшем ИСПОЛНИТЕЛЬ, "
            f"с другой стороны, составили настоящий Акт о том, что выполненные Исполнителем работы:"
        ).font.size = Pt(11)
    p.paragraph_format.space_after = Pt(0)

    doc.add_paragraph("Работы:", style="Normal").paragraph_format.space_after = Pt(0)
    for idx, work in enumerate(work_list, 1):
        doc.add_paragraph(f"1.{idx}. {work}.", style="Normal").paragraph_format.space_after = Pt(0)
    doc.add_paragraph(
        f"удовлетворяют условиям Договора № {contract_number} от {doc_date} г.",
        style="Normal",
    ).paragraph_format.space_after = Pt(0)

    rubles = int(total_cost)
    sum_words = num2words(rubles, lang="ru", to="cardinal")
    doc.add_paragraph(
        f"2. Договорная цена выполненных работ составляет {total_cost:.2f} "
        f"({sum_words.capitalize()} белорусских рублей 00 коп.). Без НДС.",
        style="Normal",
    ).paragraph_format.space_after = Pt(0)
    doc.add_paragraph("3. Исполнитель выполнил обязательства в полном объёме.", style="Normal").paragraph_format.space_after = Pt(0)
    doc.add_paragraph("4. Заказчик выполненные работы принял, претензий не имеет.", style="Normal").paragraph_format.space_after = Pt(0)
    doc.add_paragraph(
        "5. Настоящий Акт составлен в 2(двух) экземплярах, один у Исполнителя, второй — у Заказчика.",
        style="Normal",
    ).paragraph_format.space_after = Pt(0)

    is_individual = customer.get("client_type") == "individual"
    _add_requisites_and_signatures(
        doc,
        customer,
        executor_short_name=SHORT_NAME,
        customer_short_name=customer.get("short_name", customer.get("rep_name", customer.get("org_name", ""))),
        is_individual=is_individual,
    )

    save_dir = Path(save_dir or OUTPUT_DIR)
    save_dir.mkdir(parents=True, exist_ok=True)
    file_path = save_dir / f"act_{contract_number}.docx"
    doc.save(file_path)
    return str(file_path)
